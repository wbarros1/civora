"""Tests voor lokale CV-tekstextractie."""

from io import BytesIO

import pytest

from docx import Document
from pypdf import PdfWriter

from backend.app.services.candidate_cv_text import (
    DOCX_MIME_TYPE,
    PDF_MIME_TYPE,
    CvTextExtractionError,
    UnreadableCvTextError,
    UnsupportedCvTextTypeError,
    extract_cv_text,
    normalize_cv_text,
)


def create_docx_bytes() -> bytes:
    """Maak een representatief DOCX-CV in geheugen."""

    document = Document()

    document.add_paragraph(
        "Wilson Testkandidaat"
    )

    document.add_paragraph(
        "Data Engineer met ervaring "
        "in Python en SQL."
    )

    table = document.add_table(
        rows=2,
        cols=2,
    )

    table.cell(
        0,
        0,
    ).text = "Werkervaring"

    table.cell(
        0,
        1,
    ).text = (
        "Data Engineer bij "
        "Organisatie X"
    )

    table.cell(
        1,
        0,
    ).text = "Technologie"

    table.cell(
        1,
        1,
    ).text = (
        "Python, SQL, Azure"
    )

    output = BytesIO()

    document.save(
        output
    )

    return output.getvalue()


def create_blank_pdf_bytes() -> bytes:
    """Maak een geldige PDF zonder tekstlaag."""

    writer = PdfWriter()

    writer.add_blank_page(
        width=612,
        height=792,
    )

    output = BytesIO()

    writer.write(
        output
    )

    return output.getvalue()


def test_normalize_cv_text_preserves_structure() -> None:
    value = (
        "  Data   Engineer  \r\n"
        "\r\n"
        "\r\n"
        " Python\tSQL "
    )

    result = normalize_cv_text(
        value
    )

    assert result == (
        "Data Engineer\n\n"
        "Python SQL"
    )


def test_docx_extracts_paragraphs_and_tables() -> None:
    result = extract_cv_text(
        content=create_docx_bytes(),
        mime_type=DOCX_MIME_TYPE,
    )

    assert (
        result.source_type
        == "docx"
    )

    assert (
        result.page_count
        is None
    )

    assert (
        "Wilson Testkandidaat"
        in result.text
    )

    assert (
        "Data Engineer bij "
        "Organisatie X"
        in result.text
    )

    assert (
        "Python, SQL, Azure"
        in result.text
    )

    assert (
        result.character_count
        >= 50
    )


def test_blank_pdf_is_rejected_as_unreadable() -> None:
    with pytest.raises(
        UnreadableCvTextError
    ):
        extract_cv_text(
            content=(
                create_blank_pdf_bytes()
            ),
            mime_type=(
                PDF_MIME_TYPE
            ),
        )


def test_invalid_pdf_is_rejected() -> None:
    with pytest.raises(
        CvTextExtractionError
    ):
        extract_cv_text(
            content=(
                b"%PDF-1.4\n"
                b"dit is geen geldige pdf"
            ),
            mime_type=(
                PDF_MIME_TYPE
            ),
        )


def test_invalid_docx_is_rejected() -> None:
    with pytest.raises(
        CvTextExtractionError
    ):
        extract_cv_text(
            content=(
                b"dit is geen docx"
            ),
            mime_type=(
                DOCX_MIME_TYPE
            ),
        )


def test_empty_cv_is_rejected() -> None:
    with pytest.raises(
        CvTextExtractionError
    ):
        extract_cv_text(
            content=b"",
            mime_type=(
                PDF_MIME_TYPE
            ),
        )


def test_unknown_mime_type_is_rejected() -> None:
    with pytest.raises(
        UnsupportedCvTextTypeError
    ):
        extract_cv_text(
            content=(
                b"voldoende inhoud "
                * 20
            ),
            mime_type="text/plain",
        )