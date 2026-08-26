"""Lokale tekstextractie uit kandidaat-CV's."""

from __future__ import annotations

import re
import unicodedata

from dataclasses import dataclass
from io import BytesIO
from typing import Literal

from docx import Document
from docx.document import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph
from pypdf import PdfReader


PDF_MIME_TYPE = "application/pdf"

DOCX_MIME_TYPE = (
    "application/vnd.openxmlformats-officedocument."
    "wordprocessingml.document"
)

MIN_READABLE_CHARACTERS = 50


class CvTextExtractionError(
    ValueError
):
    """Basisklasse voor fouten bij lokale CV-tekstextractie."""


class UnsupportedCvTextTypeError(
    CvTextExtractionError
):
    """Het bestandstype wordt niet ondersteund."""


class EncryptedPdfError(
    CvTextExtractionError
):
    """De PDF is versleuteld en kan niet worden gelezen."""


class UnreadableCvTextError(
    CvTextExtractionError
):
    """Het CV bevat onvoldoende uitleesbare tekst."""


@dataclass(
    frozen=True
)
class ExtractedCvText:
    """Lokaal uitgelezen en genormaliseerde CV-tekst."""

    text: str

    source_type: Literal[
        "pdf",
        "docx",
    ]

    page_count: (
        int
        | None
    )

    character_count: int


def normalize_cv_text(
    value: str,
) -> str:
    """
    Normaliseer CV-tekst zonder inhoudelijk
    nieuwe informatie te introduceren.

    Regeleinden blijven behouden zodat
    structuur en evidence later bruikbaar
    blijven.
    """

    normalized_value = (
        unicodedata.normalize(
            "NFKC",
            value,
        )
        .replace(
            "\u00a0",
            " ",
        )
        .replace(
            "\r\n",
            "\n",
        )
        .replace(
            "\r",
            "\n",
        )
    )

    normalized_lines: list[str] = []

    for line in normalized_value.split(
        "\n"
    ):
        cleaned_line = re.sub(
            r"[ \t]+",
            " ",
            line,
        ).strip()

        normalized_lines.append(
            cleaned_line
        )

    normalized_value = "\n".join(
        normalized_lines
    )

    normalized_value = re.sub(
        r"\n{3,}",
        "\n\n",
        normalized_value,
    )

    return normalized_value.strip()


def count_readable_characters(
    value: str,
) -> int:
    """Tel niet-witruimtekarakters."""

    return len(
        re.sub(
            r"\s+",
            "",
            value,
        )
    )


def ensure_readable_cv_text(
    value: str,
) -> None:
    """
    Weiger lege of vrijwel lege extracties.

    Dit voorkomt dat een scan-PDF zonder
    tekstlaag verder de LLM-pipeline ingaat.
    """

    character_count = (
        count_readable_characters(
            value
        )
    )

    if (
        character_count
        < MIN_READABLE_CHARACTERS
    ):
        raise UnreadableCvTextError(
            "Het CV bevat onvoldoende "
            "uitleesbare tekst. "
            "Upload een tekst-PDF of DOCX."
        )


def extract_pdf_text(
    content: bytes,
) -> ExtractedCvText:
    """Lees tekst lokaal uit een PDF."""

    try:
        reader = PdfReader(
            BytesIO(
                content
            ),
            strict=False,
        )

    except Exception as exc:
        raise CvTextExtractionError(
            "De PDF kon niet worden gelezen."
        ) from exc

    if reader.is_encrypted:
        try:
            decrypt_result = (
                reader.decrypt(
                    ""
                )
            )

        except Exception as exc:
            raise EncryptedPdfError(
                "De PDF is beveiligd en "
                "kan niet worden gelezen."
            ) from exc

        if not decrypt_result:
            raise EncryptedPdfError(
                "De PDF is beveiligd en "
                "kan niet worden gelezen."
            )

    page_texts: list[str] = []

    try:
        for page in reader.pages:
            page_text = (
                page.extract_text()
                or ""
            )

            page_texts.append(
                page_text
            )

    except Exception as exc:
        raise CvTextExtractionError(
            "De tekst uit de PDF kon "
            "niet worden uitgelezen."
        ) from exc

    normalized_text = (
        normalize_cv_text(
            "\n\n".join(
                page_texts
            )
        )
    )

    ensure_readable_cv_text(
        normalized_text
    )

    return ExtractedCvText(
        text=normalized_text,
        source_type="pdf",
        page_count=len(
            reader.pages
        ),
        character_count=(
            count_readable_characters(
                normalized_text
            )
        ),
    )


def _iter_docx_body_text(
    document: DocxDocument,
):
    """
    Loop door paragrafen en tabellen in
    dezelfde volgorde als in het DOCX-document.
    """

    for child in (
        document
        .element
        .body
        .iterchildren()
    ):
        if child.tag.endswith(
            "}p"
        ):
            paragraph = Paragraph(
                child,
                document,
            )

            text = paragraph.text.strip()

            if text:
                yield text

        elif child.tag.endswith(
            "}tbl"
        ):
            table = Table(
                child,
                document,
            )

            for row in table.rows:
                cell_values: list[str] = []

                seen_values: set[str] = set()

                for cell in row.cells:
                    cell_text = " ".join(
                        paragraph.text.strip()
                        for paragraph
                        in cell.paragraphs
                        if paragraph.text.strip()
                    )

                    cell_text = (
                        normalize_cv_text(
                            cell_text
                        )
                    )

                    if not cell_text:
                        continue

                    comparison_value = (
                        cell_text.casefold()
                    )

                    if (
                        comparison_value
                        in seen_values
                    ):
                        continue

                    seen_values.add(
                        comparison_value
                    )

                    cell_values.append(
                        cell_text
                    )

                if cell_values:
                    yield " | ".join(
                        cell_values
                    )


def _iter_docx_header_footer_text(
    document: DocxDocument,
):
    """
    Lees eenvoudige header/footer-tekst mee.

    Contactgegevens staan bij CV's regelmatig
    in een header of footer.
    """

    seen_values: set[str] = set()

    for section in document.sections:
        for container in (
            section.header,
            section.footer,
        ):
            for paragraph in (
                container.paragraphs
            ):
                text = normalize_cv_text(
                    paragraph.text
                )

                if not text:
                    continue

                comparison_value = (
                    text.casefold()
                )

                if (
                    comparison_value
                    in seen_values
                ):
                    continue

                seen_values.add(
                    comparison_value
                )

                yield text

            for table in container.tables:
                for row in table.rows:
                    cell_values: list[str] = []

                    for cell in row.cells:
                        cell_text = (
                            " ".join(
                                paragraph.text.strip()
                                for paragraph
                                in cell.paragraphs
                                if paragraph.text.strip()
                            )
                        )

                        cell_text = (
                            normalize_cv_text(
                                cell_text
                            )
                        )

                        if cell_text:
                            cell_values.append(
                                cell_text
                            )

                    if cell_values:
                        row_text = " | ".join(
                            cell_values
                        )

                        comparison_value = (
                            row_text.casefold()
                        )

                        if (
                            comparison_value
                            in seen_values
                        ):
                            continue

                        seen_values.add(
                            comparison_value
                        )

                        yield row_text


def extract_docx_text(
    content: bytes,
) -> ExtractedCvText:
    """Lees tekst lokaal uit een DOCX."""

    try:
        document = Document(
            BytesIO(
                content
            )
        )

    except Exception as exc:
        raise CvTextExtractionError(
            "Het DOCX-bestand kon "
            "niet worden gelezen."
        ) from exc

    text_parts: list[str] = []

    text_parts.extend(
        _iter_docx_header_footer_text(
            document
        )
    )

    text_parts.extend(
        _iter_docx_body_text(
            document
        )
    )

    normalized_text = (
        normalize_cv_text(
            "\n\n".join(
                text_parts
            )
        )
    )

    ensure_readable_cv_text(
        normalized_text
    )

    return ExtractedCvText(
        text=normalized_text,
        source_type="docx",
        page_count=None,
        character_count=(
            count_readable_characters(
                normalized_text
            )
        ),
    )


def extract_cv_text(
    *,
    content: bytes,
    mime_type: str,
) -> ExtractedCvText:
    """Selecteer de juiste lokale CV-parser."""

    if not content:
        raise CvTextExtractionError(
            "Het CV-bestand is leeg."
        )

    if (
        mime_type
        == PDF_MIME_TYPE
    ):
        return extract_pdf_text(
            content
        )

    if (
        mime_type
        == DOCX_MIME_TYPE
    ):
        return extract_docx_text(
            content
        )

    raise UnsupportedCvTextTypeError(
        "Alleen PDF en DOCX worden "
        "ondersteund."
    )